"""
Builds a one-page Word summary of a wrapper test session and zips the
whole session folder (including that summary) into a single file an
engineer can hand off - the session folder itself has qxdm/, reports/,
syslog/, and metadata/ split into separate subfolders, which isn't
something you can just email or attach to a ticket as-is.
"""

from __future__ import annotations

import csv
import json
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document


def _read_wrapper_metadata(session_folder: Path) -> dict[str, Any]:
    metadata_file = (
        session_folder / "metadata" / "wrapper_session.json"
    )

    try:
        return json.loads(
            metadata_file.read_text(encoding="utf-8")
        )
    except (OSError, ValueError):
        return {}


def _read_throughput_summary(
    session_folder: Path,
) -> dict[str, Any]:
    csv_path = (
        session_folder / "reports" / "throughput_results.csv"
    )

    if not csv_path.exists():
        return {}

    def _floats(
        rows: list[dict[str, str]],
        key: str,
    ) -> list[float]:
        values: list[float] = []

        for row in rows:
            raw = (row.get(key) or "").strip()

            if not raw:
                continue

            try:
                values.append(float(raw))
            except ValueError:
                continue

        return values

    try:
        with csv_path.open(
            "r",
            encoding="utf-8-sig",
            newline="",
        ) as csv_file:
            rows = list(csv.DictReader(csv_file))
    except OSError:
        return {}

    if not rows:
        return {}

    def _avg(values: list[float]) -> float | None:
        return (
            round(sum(values) / len(values), 2)
            if values
            else None
        )

    return {
        "run_count": len(rows),
        "average_download_mbps": _avg(
            _floats(rows, "download_mbps")
        ),
        "average_upload_mbps": _avg(
            _floats(rows, "upload_mbps")
        ),
        "average_ping_ms": _avg(
            _floats(rows, "ping_ms")
        ),
    }


def _read_radio_metrics(
    session_folder: Path,
) -> dict[str, Any] | None:
    # Radio metrics aren't captured automatically as part of a wrapper
    # session today - this only picks something up if a radio_metrics.json
    # snapshot was explicitly saved into reports/. Skipped in the report
    # entirely when that file doesn't exist, rather than showing an empty
    # section.
    metrics_file = (
        session_folder / "reports" / "radio_metrics.json"
    )

    if not metrics_file.exists():
        return None

    try:
        return json.loads(
            metrics_file.read_text(encoding="utf-8")
        )
    except (OSError, ValueError):
        return None


def _list_artifact_files(
    folder: Path,
) -> list[dict[str, Any]]:
    if not folder.exists():
        return []

    files: list[dict[str, Any]] = []

    for path in sorted(folder.rglob("*")):
        if not path.is_file():
            continue

        stat = path.stat()

        files.append(
            {
                "name": str(
                    path.relative_to(folder)
                ),
                "size_bytes": stat.st_size,
                "modified_at": datetime.fromtimestamp(
                    stat.st_mtime
                ).strftime("%Y-%m-%d %H:%M:%S"),
            }
        )

    return files


def _format_size(size_bytes: int) -> str:
    if size_bytes >= 1024 * 1024:
        return f"{size_bytes / (1024 * 1024):.1f} MB"

    if size_bytes >= 1024:
        return f"{size_bytes / 1024:.1f} KB"

    return f"{size_bytes} bytes"


def generate_session_report(session_folder: Path) -> Path:
    """
    Build reports/Session_Report.docx summarizing this wrapper session:
    session info, throughput averages, radio metrics (if captured), and
    every QXDM/syslog file with its size - so the whole session can be
    understood without opening each subfolder individually.
    """
    session_folder = Path(session_folder).resolve()

    metadata = _read_wrapper_metadata(session_folder)
    throughput_summary = _read_throughput_summary(
        session_folder
    )
    radio_metrics = _read_radio_metrics(session_folder)
    qxdm_files = _list_artifact_files(
        session_folder / "qxdm"
    )
    syslog_files = _list_artifact_files(
        session_folder / "syslog"
    )

    document = Document()

    document.add_heading(
        "WNC TestHub Session Report",
        level=1,
    )

    document.add_paragraph(
        "Generated "
        + datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    )

    def add_two_column_table(
        heading: str,
        rows: list[tuple[str, Any]],
    ) -> None:
        document.add_heading(heading, level=2)

        table = document.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"

        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = (
                str(value)
                if value not in (None, "")
                else "Not available"
            )

    add_two_column_table(
        "Session Information",
        [
            (
                "Session Name",
                metadata.get(
                    "session_name", session_folder.name
                ),
            ),
            ("Titan IP", metadata.get("titan_ip")),
            ("Mode", metadata.get("mode")),
            ("Created At", metadata.get("created_at")),
            ("Session Folder", str(session_folder)),
        ],
    )

    if throughput_summary:
        add_two_column_table(
            "Throughput Summary",
            [
                (
                    "Test Runs",
                    throughput_summary.get("run_count"),
                ),
                (
                    "Average Download",
                    (
                        f"{throughput_summary['average_download_mbps']} Mbps"
                        if throughput_summary.get(
                            "average_download_mbps"
                        )
                        is not None
                        else None
                    ),
                ),
                (
                    "Average Upload",
                    (
                        f"{throughput_summary['average_upload_mbps']} Mbps"
                        if throughput_summary.get(
                            "average_upload_mbps"
                        )
                        is not None
                        else None
                    ),
                ),
                (
                    "Average Ping",
                    (
                        f"{throughput_summary['average_ping_ms']} ms"
                        if throughput_summary.get(
                            "average_ping_ms"
                        )
                        is not None
                        else None
                    ),
                ),
            ],
        )
    else:
        document.add_heading(
            "Throughput Summary", level=2
        )
        document.add_paragraph(
            "No throughput results were recorded for this session."
        )

    if radio_metrics:
        add_two_column_table(
            "Radio Metrics",
            [
                (
                    "Technology",
                    radio_metrics.get("technology"),
                ),
                ("Mode", radio_metrics.get("mode")),
                (
                    "Serving Band",
                    radio_metrics.get("serving_band"),
                ),
                (
                    "RSRP",
                    (
                        f"{radio_metrics['rsrp_dbm']} dBm"
                        if radio_metrics.get("rsrp_dbm")
                        is not None
                        else None
                    ),
                ),
                (
                    "RSSI",
                    (
                        f"{radio_metrics['rssi_dbm']} dBm"
                        if radio_metrics.get("rssi_dbm")
                        is not None
                        else None
                    ),
                ),
                (
                    "SINR",
                    (
                        f"{radio_metrics['sinr_db']} dB"
                        if radio_metrics.get("sinr_db")
                        is not None
                        else None
                    ),
                ),
                (
                    "Firmware",
                    radio_metrics.get("firmware_version"),
                ),
            ],
        )

    def add_file_section(
        heading: str,
        files: list[dict[str, Any]],
    ) -> None:
        document.add_heading(heading, level=2)

        if not files:
            document.add_paragraph(
                "No files were found."
            )
            return

        table = document.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"

        header_cells = table.rows[0].cells
        header_cells[0].text = "File"
        header_cells[1].text = "Size"
        header_cells[2].text = "Modified"

        for file_info in files:
            row_cells = table.add_row().cells
            row_cells[0].text = file_info["name"]
            row_cells[1].text = _format_size(
                file_info["size_bytes"]
            )
            row_cells[2].text = file_info["modified_at"]

    add_file_section("QXDM Log Files", qxdm_files)
    add_file_section("Syslog Files", syslog_files)

    reports_folder = session_folder / "reports"
    reports_folder.mkdir(parents=True, exist_ok=True)

    report_path = reports_folder / "Session_Report.docx"
    document.save(report_path)

    return report_path


def zip_session_folder(
    session_folder: Path,
    include_report: bool = True,
) -> Path:
    """
    Zip the whole wrapper session folder (qxdm/, reports/, syslog/,
    metadata/, and anything else in it) into one file next to it.
    """
    session_folder = Path(session_folder).resolve()

    if include_report:
        generate_session_report(session_folder)

    zip_path = (
        session_folder.parent / f"{session_folder.name}.zip"
    )

    # Rebuild from scratch each time rather than appending - avoids
    # stale duplicate entries if this is run more than once for the
    # same session.
    if zip_path.exists():
        zip_path.unlink()

    with zipfile.ZipFile(
        zip_path, "w", zipfile.ZIP_DEFLATED
    ) as zip_file:
        for path in session_folder.rglob("*"):
            if path.is_file():
                zip_file.write(
                    path,
                    path.relative_to(session_folder.parent),
                )

    return zip_path
